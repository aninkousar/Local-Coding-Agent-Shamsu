from __future__ import annotations
from pathlib import Path

IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif"}
TEXT_EXTS = {".txt", ".md", ".markdown", ".rst", ".json", ".yaml", ".yml", ".csv"}


class UnsupportedFile(Exception):
    pass


def classify(path: Path) -> str:
    ext = path.suffix.lower()
    if ext in IMAGE_EXTS:
        return "image"
    if ext == ".pdf":
        return "pdf"
    if ext == ".docx":
        return "docx"
    if ext in TEXT_EXTS or ext == "":
        return "text"
    return "text"  # best effort: try to read as text/code


def read_as_text(path: Path) -> str:
    """Extract plain text content from a document. Images are NOT handled here -
    use image_to_b64 + pass to the vision model instead (see ollama_client).
    """
    kind = classify(path)

    if kind == "text":
        return path.read_text(encoding="utf-8", errors="replace")

    if kind == "pdf":
        try:
            from pypdf import PdfReader
        except ImportError:
            raise UnsupportedFile(
                "Reading PDFs needs an extra package. Run: pip install pypdf"
            )
        reader = PdfReader(str(path))
        pages = [p.extract_text() or "" for p in reader.pages]
        text = "\n\n".join(pages).strip()
        if not text:
            raise UnsupportedFile(
                f"{path} looks like a scanned/image-only PDF with no extractable text. "
                "Try exporting the relevant pages as PNG/JPG and pass those instead - "
                "the model can read images directly."
            )
        return text

    if kind == "docx":
        try:
            import docx
        except ImportError:
            raise UnsupportedFile(
                "Reading Word docs needs an extra package. Run: pip install python-docx"
            )
        d = docx.Document(str(path))
        parts = [p.text for p in d.paragraphs]
        for table in d.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return "\n".join(parts)

    raise UnsupportedFile(f"Don't know how to read {path}")
